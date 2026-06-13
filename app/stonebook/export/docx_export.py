"""DOCX-Analysebericht pro Objekt (Nachfolger des Template-v4_2-Workflows)."""
from pathlib import Path

from docx import Document
from docx.shared import Cm

from stonebook.fields import CATEGORY_LABELS, DATA_FIELDS, FIELD_GROUPS, IMAGE_CATEGORIES, is_empty
from stonebook.migration.id_utils import display_name

MAX_IMG_WIDTH_CM = 12.0


def export_docx(conn, root: Path, obj_id: str, out_path: Path | None = None) -> Path:
    row = conn.execute("SELECT * FROM objects WHERE obj_id = ?", (obj_id,)).fetchone()
    if row is None:
        raise ValueError(f"Objekt {obj_id} nicht gefunden")
    images = conn.execute(
        "SELECT * FROM images WHERE obj_id = ? ORDER BY kategorie, dateiname",
        (obj_id,)).fetchall()
    aliases = [r[0] for r in conn.execute(
        "SELECT alias_id FROM aliases WHERE canonical_id = ? ORDER BY alias_id",
        (obj_id,)).fetchall()]

    doc = Document()
    doc.add_heading(f"StoneBoock – Analysebericht {display_name(obj_id)}", level=0)
    sub = doc.add_paragraph(f"Objekt-ID: {obj_id}")
    if row["Name"]:
        sub.add_run(f"  ·  {row['Name']}")
    if aliases:
        doc.add_paragraph("Enthält gemergte Objekte: "
                          + ", ".join(display_name(a) for a in aliases))

    doc.add_heading("1. Objektdaten", level=1)
    for group in FIELD_GROUPS:
        entries = []
        for fdef in DATA_FIELDS:
            if fdef.group != group:
                continue
            v = row[fdef.name]
            if is_empty(v):
                continue
            entries.append((fdef.label, str(v)))
        if not entries:
            continue
        doc.add_heading(group, level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for label, value in entries:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
    if row["notizen"]:
        doc.add_heading("Notizen", level=2)
        doc.add_paragraph(row["notizen"])

    doc.add_heading("2. Medien (Fotos/UV)", level=1)
    if not images:
        doc.add_paragraph("Keine Bilder vorhanden.")
    for cat in IMAGE_CATEGORIES:
        cat_rows = [r for r in images if r["kategorie"] == cat]
        if not cat_rows:
            continue
        doc.add_heading(CATEGORY_LABELS[cat], level=2)
        for r in cat_rows:
            img_path = root / r["rel_path"]
            if not img_path.is_file():
                doc.add_paragraph(f"[fehlend] {r['rel_path']}")
                continue
            try:
                doc.add_picture(str(img_path), width=Cm(MAX_IMG_WIDTH_CM))
                cap = doc.add_paragraph(r["dateiname"] or img_path.name)
                if r["herkunft_obj_id"]:
                    cap.add_run(f"  (ursprünglich {r['herkunft_obj_id']})")
            except Exception:
                doc.add_paragraph(f"[nicht einbettbar] {r['rel_path']}")

    if out_path is None:
        out_path = _default_report_path(root, obj_id)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _default_report_path(root: Path, obj_id: str) -> Path:
    num = int(obj_id.split("_")[1])
    return root / "objects" / obj_id / f"Objekt_{num:03d}_Analysebericht.docx"


def export_docx_batch(conn, root: Path, obj_ids: list[str] | None = None,
                      out_dir: Path | None = None, progress=None) -> list[Path]:
    """Schreibt Analyseberichte fuer mehrere Objekte.

    ``obj_ids=None`` exportiert alle aktiven Objekte. ``out_dir`` legt alle
    Berichte in denselben Ordner (Dateiname ``Objekt_NNN_Analysebericht.docx``);
    ohne ``out_dir`` landet jeder Bericht unter ``objects/<obj_id>/`` wie beim
    Einzelexport. ``progress(done, total, obj_id)`` wird optional pro Objekt
    aufgerufen.
    """
    if obj_ids is None:
        obj_ids = [r[0] for r in conn.execute(
            "SELECT obj_id FROM objects WHERE status='aktiv' ORDER BY obj_id"
        ).fetchall()]
    if out_dir is not None:
        out_dir.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    total = len(obj_ids)
    for i, obj_id in enumerate(obj_ids, 1):
        target = (out_dir / _default_report_path(root, obj_id).name
                  if out_dir is not None else None)
        written.append(export_docx(conn, root, obj_id, target))
        if progress is not None:
            progress(i, total, obj_id)
    return written

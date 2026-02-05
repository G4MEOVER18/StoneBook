import pandas as pd
from docx import Document

# Eingabedateien
docx_input = "Objekt_043_Analysebericht.docx"
csv_file = "stoneboock_daten_v2.csv"
template_file = "StoneBoock_SingleObject_Template_v2.docx"

# 1. Inhalte aus Analyse-DOCX lesen
doc_in = Document(docx_input)
full_text = "\n".join([p.text for p in doc_in.paragraphs])

# 2. Neues DOCX aus Template bauen
doc_out = Document(template_file)
doc_out.add_heading("Analysebericht – Objekt 043", level=1)
doc_out.add_paragraph(full_text)

docx_output = "Analyse_Objekt_043_Template.docx"
doc_out.save(docx_output)
print(f"[+] Neuer Bericht gespeichert: {docx_output}")

# 3. CSV ergänzen
df = pd.read_csv(csv_file)

new_row = {
    "ID": "OBJ-043",
    "Name": "Objekt 43",
    "Fundort": "Aach, Amriswil TG",
    "Beschreibung": "Weißer Quarz mit braun-rötlichen Eisenoxidadern, blockig, grobkörnig",
    "Mineralart": "Quarz mit Eisenoxiden",
    "Gesteinsart": "Gangquarz",
    "Kristallsystem": "Trigonal",
    "Mohs_Haerte_min": 7,
    "Mohs_Haerte_max": 7,
    "Dichte_min_gcm3": 2.65,
    "Dichte_max_gcm3": 2.65,
    "Spaltbarkeit": "keine",
    "Bruch": "muschelig-splittrig",
    "Glanz": "glasig, matt-erdig",
    "Transparenz": "opak – transluzent",
    "Farbe_beobachtet": "weiß mit braun-rötlichen Adern",
    "Strichfarbe": "weiß (Quarz) / rotbraun (Fe-Oxid)",
    "UV_365nm": "keine",
    "UV_254nm": "keine",
    "Magnetismus": "nein",
    "HCl_Reaktion": "keine",
    "Reaktionshinweis": "Quarz reagiert nicht, Fe-Oxide überdecken",
    "Gewicht_g": 41.0,
    "Laenge_mm": 62,
    "Breite_mm": 47,
    "Hoehe_mm": 28,
    "Seltenheit_global_1_10": 2,
    "Seltenheit_Fundort_1_10": 4,
    "Nachfrage_1_10": 3,
    "Wert_CHF_roh": "gering",
    "Wert_CHF_poliert": "gering",
    "Wert_CHF_Schmuck": "keiner",
    "Wert_USD_Talisman": "gering",
    "Marktwert_Industrie": "gering",
    "Wissenschaftlicher_Wert_CHF": "mittel",
    "Beste_Verwendung": "Sammlerstück, Dekoration",
    "Pruefempfehlungen": "Ritztest, HCl-Test, Strichprobe",
    "Confidence_Prozent": 90
}

# Alte Zeile ersetzen, falls vorhanden
df = df[df["ID"] != "OBJ-043"]
df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
df.to_csv(csv_file, index=False, encoding="utf-8-sig")
print(f"[+] CSV aktualisiert: {csv_file}")
